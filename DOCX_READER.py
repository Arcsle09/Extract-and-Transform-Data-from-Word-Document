# -*- coding: utf-8 -*-
"""
Created on Sat Dec 18 15:58:04 2021

@author: chra8017
"""

from docx import Document

import pandas as pd

output = pd.DataFrame(columns = ['Company Name','Rest of the Details'])

document = Document('5_6190415455547557036-converted.docx')

total_lines = len(document.paragraphs)

all_data = document.paragraphs

#removing first 2 spaces

all_data.pop(0)

#Keep Track of new company Name
new_company_tracker = 0

#Kepe track of each statements
para_tracker = 0

#max_para_num = 10000

#collecting company name and its rest of the details.
for para_iterator in range(1,len(all_data)):
    company_details = []
    if (all_data[para_iterator].text == '') and (all_data[para_iterator-1].text != '') and (all_data[para_iterator-1].text != ''):
        
        for detail_index in range(para_tracker+1,para_iterator):
            company_details.append(all_data[detail_index].text)
        output.loc[new_company_tracker,'Company Name'] = company_details[0]
        output.loc[new_company_tracker,'Rest of the Details'] = company_details[1:]
        para_tracker = para_iterator
        new_company_tracker = new_company_tracker + 1
        
    elif para_iterator == len(all_data)-1:
        for detail_index in range(para_tracker+1,para_iterator):
            company_details.append(all_data[detail_index].text)
        output.loc[new_company_tracker,'Company Name'] = company_details[0]
        output.loc[new_company_tracker,'Rest of the Details'] = company_details[1:]

#Merging the rows if the rest of the details got spilled to next rows.
for df_iterator in range(0,len(output)):
    if (output.loc[df_iterator,'Company Name'] == ''):
        merged_data_at_new_page = output.loc[df_iterator-1,'Rest of the Details'] + output.loc[df_iterator,'Rest of the Details']
        output.loc[df_iterator-1,'Rest of the Details'] = [merge for merge in merged_data_at_new_page if merge !='']

filtered_output = output[output['Company Name'] != '']
filtered_output.reset_index(inplace=True)
filtered_output.pop('index')

for df_iterator in range(0,len(filtered_output)):
    if "imdb" in filtered_output.loc[df_iterator,'Company Name']:
        merged_data_at_new_page = filtered_output.loc[df_iterator-1,'Rest of the Details'] + [filtered_output.loc[df_iterator,'Company Name']] + filtered_output.loc[df_iterator,'Rest of the Details']
        filtered_output.loc[df_iterator-1,'Rest of the Details'] = [merge for merge in merged_data_at_new_page if merge !='']

final_output = filtered_output[~filtered_output['Company Name'].str.contains('imdb')]

#replacing "\n" from company name column
final_output['Company Name'] = final_output['Company Name'].str.replace("\n","")

final_output.reset_index(inplace=True)
final_output.pop('index')

#collecting the index of unnecessary rows and then remove them.
rows_index_to_remove = []

for df_iter in range(0,len(final_output)):
    #final_output.loc[df_iterator,'Company Name'][0].isdigit()==False) and (
    if final_output.loc[df_iter,'Company Name'].isupper()==False:
        if not final_output.loc[df_iter,'Company Name'][0].isdigit():
            merged_data_at_new_page = final_output.loc[df_iter-1,'Rest of the Details'] + [final_output.loc[df_iter,'Company Name']] + final_output.loc[df_iter,'Rest of the Details']
            final_output.loc[df_iter-1,'Rest of the Details'] = [merge for merge in merged_data_at_new_page if merge !='']
            rows_index_to_remove.append(df_iter)

final_output.drop(final_output.index[rows_index_to_remove],inplace=True)
        
final_output.reset_index(inplace=True)
final_output.pop('index')            
        
final_output.to_csv('raw_output_all.csv',index=False,encoding='utf-8-sig')

# layout_file = pd.DataFrame(columns = ['Company Name','Company Type','Address and Phone Number','Website','IMDB PAGE','Submission Policy','Project Type','Genres'])

# for filter_iter in range(0,len(final_output)):
    
#     layout_file.loc[filter_iter,'Company Name'] = final_output.loc[filter_iter,'Company Name']
    
#     layout_file.loc[filter_iter,'Company Type'] = final_output.loc[filter_iter,'Rest of the Details'][0]
    
#     Address_phone_extract = final_output.loc[filter_iter,'Rest of the Details'][1]
    
#     #Address_Plus_Phone = " ".join([add_phone for add_phone in Address_phone_extract if "\n" not in add_phone])
    
#     layout_file.loc[filter_iter,'Address and Phone Number'] = Address_phone_extract.replace("\n","")
    
    
